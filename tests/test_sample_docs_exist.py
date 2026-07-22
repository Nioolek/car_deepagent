import re

from docx import Document

from car_deepagent.paths import interviews_dir


def test_three_interviews_exist():
    directory = interviews_dir()
    for name in ("interview_001.docx", "interview_002.docx", "interview_003.docx"):
        assert (directory / name).exists(), name


def test_interviews_match_mock_profiles_and_include_disclaimer():
    expected = {
        "interview_001.docx": ("U001", "陈思远", "上海", "问界 M7"),
        "interview_002.docx": ("U002", "林婉清", "深圳", "智界 R7"),
        "interview_003.docx": ("U003", "周启明", "杭州", "享界 S9"),
    }
    for filename, profile_values in expected.items():
        document = Document(interviews_dir() / filename)
        text = "\n".join(paragraph.text for paragraph in document.paragraphs)
        assert "虚构样例，非真实用户数据" in text
        for value in profile_values:
            assert value in text


def test_first_interview_has_at_least_20000_chinese_characters():
    document = Document(interviews_dir() / "interview_001.docx")
    text = "\n".join(paragraph.text for paragraph in document.paragraphs)
    chinese_chars = re.findall(r"[\u3400-\u4dbf\u4e00-\u9fff]", text)
    assert len(chinese_chars) >= 20_000
