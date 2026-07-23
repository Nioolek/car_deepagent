import json
from pathlib import Path

from docx import Document

from car_deepagent.tools import documents as docs

SAMPLE_MD = """# 一、背景
用户在上海工作。

# 二、智驾体验
高速 NOA 体验良好，但雨天会接管。

# 三、总结
整体满意。
"""


def test_split_chapters():
    chapters = docs.split_chapters(SAMPLE_MD)
    assert len(chapters) == 3
    assert chapters[0]["chapter_id"] == "1"
    assert "智驾" in chapters[1]["title"]


def test_ensure_summary_tree_builds_once(tmp_path, monkeypatch):
    monkeypatch.setattr(docs, "markdown_cache_dir", lambda: tmp_path / "md")
    monkeypatch.setattr(docs, "summary_trees_dir", lambda: tmp_path / "trees")
    (tmp_path / "md").mkdir()
    (tmp_path / "trees").mkdir()
    (tmp_path / "md" / "interview_t.md").write_text(SAMPLE_MD, encoding="utf-8")

    calls = {"n": 0}

    class FakeMsg:
        content = "章节摘要"

    class FakeModel:
        def invoke(self, messages):
            calls["n"] += 1
            return FakeMsg()

    monkeypatch.setattr(docs, "build_chat_model", lambda: FakeModel())

    raw1 = docs.ensure_summary_tree.invoke({"doc_id": "interview_t"})
    raw2 = docs.ensure_summary_tree.invoke({"doc_id": "interview_t"})
    d1 = json.loads(raw1)
    d2 = json.loads(raw2)
    assert d1["cached"] is False
    assert d2["cached"] is True
    assert calls["n"] == 3

    summary = json.loads(
        docs.get_chapter_summary.invoke(
            {"doc_id": "interview_t", "chapter_id": "2"}
        )
    )
    assert summary["summary"] == "章节摘要"
    # Citation evidence comes from line-based read_file on the markdown cache.
    markdown = (tmp_path / "md" / "interview_t.md").read_text(encoding="utf-8")
    assert "NOA" in markdown or "雨天" in markdown


def test_ensure_summary_tree_does_not_cache_model_failure(tmp_path, monkeypatch):
    monkeypatch.setattr(docs, "markdown_cache_dir", lambda: tmp_path / "md")
    monkeypatch.setattr(docs, "summary_trees_dir", lambda: tmp_path / "trees")
    (tmp_path / "md").mkdir()
    (tmp_path / "trees").mkdir()
    (tmp_path / "md" / "interview_t.md").write_text(SAMPLE_MD, encoding="utf-8")

    class FailingModel:
        def invoke(self, messages):
            raise TimeoutError("summary timed out")

    monkeypatch.setattr(docs, "build_chat_model", lambda: FailingModel())

    result = json.loads(
        docs.ensure_summary_tree.invoke({"doc_id": "interview_t"})
    )

    assert "error" in result
    assert "summary timed out" in result["error"]
    assert result["chapter_id"] == "1"
    assert not (tmp_path / "trees" / "interview_t.json").exists()


def test_source_change_rebuilds_markdown_and_summary_tree(tmp_path, monkeypatch):
    monkeypatch.setattr(docs, "markdown_cache_dir", lambda: tmp_path / "md")
    monkeypatch.setattr(docs, "summary_trees_dir", lambda: tmp_path / "trees")
    (tmp_path / "md").mkdir()
    (tmp_path / "trees").mkdir()
    src = tmp_path / "interview_t.docx"

    def write_docx(body: str) -> None:
        document = Document()
        document.add_heading("背景", level=1)
        document.add_paragraph(body)
        document.save(src)

    class FakeMsg:
        content = "章节摘要"

    class FakeModel:
        def __init__(self):
            self.calls = 0

        def invoke(self, messages):
            self.calls += 1
            return FakeMsg()

    model = FakeModel()
    monkeypatch.setattr(docs, "build_chat_model", lambda: model)

    write_docx("第一版内容")
    markdown_1 = json.loads(
        docs.ensure_document_markdown.invoke({"path": str(src)})
    )
    tree_1 = json.loads(
        docs.ensure_summary_tree.invoke({"doc_id": "interview_t"})
    )

    write_docx("第二版完全不同的内容")
    markdown_2 = json.loads(
        docs.ensure_document_markdown.invoke({"path": str(src)})
    )
    tree_2 = json.loads(
        docs.ensure_summary_tree.invoke({"doc_id": "interview_t"})
    )

    assert Path(markdown_2["markdown_path"]).read_text(encoding="utf-8").find(
        "第二版完全不同的内容"
    ) >= 0
    assert markdown_1["source_sha256"] != markdown_2["source_sha256"]
    assert tree_1["markdown_sha256"] != tree_2["markdown_sha256"]
    assert tree_2["cached"] is False
    assert model.calls == 2


def test_invalid_doc_id_rejected(tmp_path, monkeypatch):
    monkeypatch.setattr(docs, "markdown_cache_dir", lambda: tmp_path / "md")
    monkeypatch.setattr(docs, "summary_trees_dir", lambda: tmp_path / "trees")
    (tmp_path / "md").mkdir()
    (tmp_path / "trees").mkdir()

    invalid = "../etc/passwd"
    for tool, args in (
        (docs.ensure_summary_tree, {"doc_id": invalid}),
        (docs.get_chapter_summary, {"doc_id": invalid, "chapter_id": "1"}),
    ):
        data = json.loads(tool.invoke(args))
        assert "error" in data
        assert "Invalid doc_id" in data["error"]

    assert list((tmp_path / "md").iterdir()) == []
    assert list((tmp_path / "trees").iterdir()) == []
    assert not (tmp_path.parent / "etc").exists()
