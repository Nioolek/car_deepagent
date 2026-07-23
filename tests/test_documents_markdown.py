import json
from pathlib import Path

from docx import Document

from car_deepagent.tools import documents as documents_mod
from car_deepagent.tools.documents import doc_id_for_path, ensure_document_markdown


def _write_docx(path: Path, paragraphs: list[str]) -> None:
    doc = Document()
    for p in paragraphs:
        doc.add_paragraph(p)
    doc.save(path)


def test_doc_id_for_path():
    assert doc_id_for_path("/tmp/interview_001.docx") == "interview_001"


def test_ensure_document_markdown_creates_cache(tmp_path, monkeypatch):
    interviews = tmp_path / "docs" / "interviews"
    interviews.mkdir(parents=True)
    cache = tmp_path / "md"
    cache.mkdir()
    monkeypatch.setattr(documents_mod, "markdown_cache_dir", lambda: cache)
    monkeypatch.setattr(
        "car_deepagent.analysis_docs.interviews_dir",
        lambda: interviews,
    )
    monkeypatch.setattr(
        "car_deepagent.analysis_docs.repo_root",
        lambda: tmp_path,
    )
    src = interviews / "interview_x.docx"
    _write_docx(src, ["背景", "受访者来自上海。", "智驾体验", "NOA 总体可用。"])
    raw = ensure_document_markdown.invoke({"path": "interview_x"})
    data = json.loads(raw)
    assert data["doc_id"] == "interview_x"
    md_path = Path(data["markdown_path"])
    assert md_path.exists()
    text = md_path.read_text(encoding="utf-8")
    assert "智驾体验" in text
    assert data["chars"] > 0


def test_ensure_document_markdown_missing_file():
    raw = ensure_document_markdown.invoke({"path": "/no/such/file.docx"})
    data = json.loads(raw)
    assert "error" in data


def test_source_change_rebuilds_markdown(tmp_path, monkeypatch):
    interviews = tmp_path / "docs" / "interviews"
    interviews.mkdir(parents=True)
    monkeypatch.setattr(documents_mod, "markdown_cache_dir", lambda: tmp_path / "md")
    monkeypatch.setattr(
        "car_deepagent.analysis_docs.interviews_dir",
        lambda: interviews,
    )
    monkeypatch.setattr(
        "car_deepagent.analysis_docs.repo_root",
        lambda: tmp_path,
    )
    (tmp_path / "md").mkdir()
    src = interviews / "interview_t.docx"

    def write_docx(body: str) -> None:
        document = Document()
        document.add_heading("背景", level=1)
        document.add_paragraph(body)
        document.save(src)

    write_docx("第一版内容")
    markdown_1 = json.loads(
        ensure_document_markdown.invoke({"path": "interview_t"})
    )
    write_docx("第二版完全不同的内容")
    markdown_2 = json.loads(
        ensure_document_markdown.invoke({"path": "interview_t"})
    )

    assert Path(markdown_2["markdown_path"]).read_text(encoding="utf-8").find(
        "第二版完全不同的内容"
    ) >= 0
    assert markdown_1["source_sha256"] != markdown_2["source_sha256"]
