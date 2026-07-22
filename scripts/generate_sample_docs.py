"""Generate deterministic, fabricated Chinese automotive interview samples."""

from __future__ import annotations

import re
from itertools import cycle
from pathlib import Path

from docx import Document


OUTPUT_DIR = Path(__file__).resolve().parents[1] / "docs" / "interviews"
DISCLAIMER = "虚构样例，非真实用户数据"
SECTIONS = {
    "背景": [
        "请先介绍一下你的日常用车场景和家庭情况。",
        "你过去开过哪些车，对智能汽车的期待是怎么形成的？",
        "工作日和周末的出行半径分别怎样？",
        "家人对车辆的关注点与你有什么不同？",
    ],
    "购车旅程": [
        "这次换车最初的触发点是什么？",
        "你对比过哪些鸿蒙智行产品，筛选标准如何变化？",
        "到店试驾和线上研究分别影响了哪些判断？",
        "最终下定前有哪些犹豫，什么信息帮助你做出决定？",
    ],
    "座舱与HMI": [
        "你最常用的座舱功能是什么，操作路径顺手吗？",
        "语音助手在多人乘坐时表现如何？",
        "车机导航、音乐和手机互联有哪些亮点或阻碍？",
        "屏幕、实体控制和场景模式之间的分工符合预期吗？",
    ],
    "智驾/NOA": [
        "你在哪些道路上使用过NOA，通常如何接管？",
        "系统的跟车、变道和上下匝道表现给你什么感受？",
        "哪些提示能建立信任，哪些行为会让你紧张？",
        "如果下一版只能改进一个智驾问题，你会选什么？",
    ],
    "服务与OTA": [
        "交付、售后和补能服务中，哪个环节印象最深？",
        "你如何看待OTA频率、更新说明和学习成本？",
        "遇到问题时更愿意使用门店、热线还是车机反馈？",
        "品牌社区和用户运营对你的长期满意度有影响吗？",
    ],
    "总结": [
        "如果向朋友推荐这辆车，你会先讲优点还是限制？",
        "目前整体满意度如何，最希望后续兑现什么承诺？",
        "这辆车对你的出行习惯带来了哪些长期变化？",
        "请用一句话概括你心中的理想智能汽车。",
    ],
}

PARTICIPANTS = {
    "001": {
        "user_id": "U001",
        "name": "陈思远",
        "profile": "三十五岁产品经理，与伴侣和孩子居住在上海，工作日跨区通勤，周末常去近郊。",
        "product": "问界M7",
        "priority": "家庭舒适、车机易用和长途辅助驾驶",
    },
    "002": {
        "user_id": "U002",
        "name": "林婉清",
        "profile": "三十岁品牌策划，独居杭州，日常城市通勤，也会带父母进行短途旅行。",
        "product": "智界S7",
        "priority": "设计质感、停车便利和跨设备体验",
    },
    "003": {
        "user_id": "U003",
        "name": "周启明",
        "profile": "四十二岁小企业经营者，与家人居住在成都，商务出行和家庭自驾频率都较高。",
        "product": "享界S9",
        "priority": "乘坐品质、服务效率和高速NOA稳定性",
    },
}

OBSERVATIONS = [
    "我不会因为一次演示就下结论，而是连续使用几周再判断。顺畅时它确实能减少操作，但边界不清楚时我会马上接管。",
    "对我来说功能数量不是重点，关键是入口稳定、反馈明确，并且家里不熟悉科技产品的人也能独立完成操作。",
    "我会把体验拆成效率、舒适和信任三个方面。效率决定每天愿不愿意用，舒适影响家人评价，信任决定长途是否敢依赖。",
    "实际道路比宣传场景复杂，施工、雨天和临时改道都可能出现。我希望系统承认不确定性，提前告诉驾驶员该做什么。",
    "好的体验往往没有存在感，第一次设置后就能自然融入行程。反复确认、层级过深或提示太晚都会打断注意力。",
    "家人的反馈会显著改变我的判断。驾驶员关注控制感，后排更在意晕车、噪声和空调，所以评价不能只看主驾。",
    "我愿意学习新功能，但更新说明需要讲清适用条件和变化原因。只列功能名称，会让我在真正上路时不敢尝试。",
    "有些问题单独看很小，连续发生就会消耗耐心。品牌如果能记录上下文并一次解决，比单纯赠送权益更能建立信任。",
]

DETAILS = [
    "早高峰拥堵时，我会观察系统是否保持合适车距，也会留意加减速会不会让后排不舒服。",
    "周末带家人出行时，老人和孩子的反应比参数更真实，他们能否顺利调节温度和音乐很重要。",
    "进入陌生商场停车场时，我更需要清晰、及时的引导，不希望在路口临时判断应该走哪条车道。",
    "夜间或下雨时，我会主动降低依赖程度，同时关注界面有没有把道路风险和系统能力边界表达清楚。",
    "长途行驶中，稳定比偶尔惊艳更重要；如果行为可预期，我能更轻松地监督系统并保留接管余量。",
    "我也会比较升级前后的变化，希望常用设置不要被重置，新的交互最好提供短教程并允许稍后查看。",
]


def chinese_character_count(text: str) -> int:
    return len(re.findall(r"[\u3400-\u4dbf\u4e00-\u9fff]", text))


def answer(profile: dict[str, str], section: str, round_number: int) -> str:
    observation = OBSERVATIONS[(round_number + len(section)) % len(OBSERVATIONS)]
    detail = DETAILS[(round_number * 2 + len(profile["name"])) % len(DETAILS)]
    return (
        f"受访者：结合我使用{profile['product']}的经历，{profile['profile']}"
        f"我主要关注{profile['priority']}。{observation}{detail}"
        f"第{round_number + 1}次回顾这个话题时，我更看重真实场景中的一致性，"
        "也希望产品团队说明设计取舍，而不是只给出理想条件下的结果。"
    )


def build_document(doc_id: str, profile: dict[str, str], target_chars: int) -> tuple[Document, int]:
    document = Document()
    document.add_heading(f"鸿蒙智行用户访谈 {doc_id}", level=0)
    document.add_paragraph(DISCLAIMER)
    document.add_paragraph(
        f"受访者编号：{profile['user_id']}　姓名：{profile['name']}　当前产品：{profile['product']}"
    )

    body_count = 0
    question_number = 1
    section_target = (target_chars + len(SECTIONS) - 1) // len(SECTIONS)
    for section, questions in SECTIONS.items():
        document.add_heading(section, level=1)
        prompts = cycle(questions)
        section_count = 0
        round_number = 0
        while section_count < section_target:
            question = next(prompts)
            response = answer(profile, section, round_number)
            paragraph = f"问{question_number}：{question}\n{response}"
            document.add_paragraph(paragraph)
            paragraph_count = chinese_character_count(paragraph)
            section_count += paragraph_count
            body_count += paragraph_count
            question_number += 1
            round_number += 1

    return document, body_count


def main() -> None:
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    targets = {"001": 20_500, "002": 5_000, "003": 5_000}
    for doc_id, profile in PARTICIPANTS.items():
        document, count = build_document(doc_id, profile, targets[doc_id])
        output_path = OUTPUT_DIR / f"interview_{doc_id}.docx"
        document.save(output_path)
        print(f"{output_path}: {count} Chinese body characters")


if __name__ == "__main__":
    main()
